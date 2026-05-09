"""
Audit Anomaly Explainer — Streamlit App (v2)
==============================================
Accepts ANY journal entry CSV, auto-detects anomalies, retrieves policies
via RAG, and generates structured audit observations using Gemini.
"""

import os
import time
import streamlit as st
import pandas as pd
import openpyxl  # noqa: F401  — needed for pd.read_excel
from pathlib import Path

from rag_pipeline import build_or_load_index
from llm_client import generate_observation, generate_observation_baseline
from anomaly_detector import run_all_detections, normalize_columns, parse_coa

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
APP_DIR = Path(__file__).parent
POLICY_DIR = APP_DIR / "policies"
INDEX_DIR = APP_DIR / "index_cache"
DATA_DIR = APP_DIR / "data"

AVAILABLE_MODELS = {
    "gemini-2.5-flash": "Gemini 2.5 Flash  ·  fast & cost-effective",
    "gemini-2.5-pro": "Gemini 2.5 Pro  ·  higher quality",
}

# Plain-language explanations for each anomaly type (manager-friendly)
FLAG_EXPLANATIONS = {
    "Duplicate Entry": {
        "icon": "📑",
        "short": "Possible Duplicate",
        "desc": "Two or more entries share the same date, amount, and accounts — this may indicate an accidental double-booking or a copy-paste error.",
    },
    "Post-Close Entry": {
        "icon": "🔒",
        "short": "Posted After Books Closed",
        "desc": "This entry was recorded after the accounting period was officially closed. Only authorized personnel should post adjustments after close.",
    },
    "Round-Dollar Amount": {
        "icon": "🎯",
        "short": "Suspiciously Round Amount",
        "desc": "The transaction amount is an exact round number (e.g. $50,000). Real invoices and expenses usually have cents — round figures may signal an estimate or fabricated entry.",
    },
    "Unusual Account Pairing": {
        "icon": "🔀",
        "short": "Unusual Account Combination",
        "desc": "The debit and credit accounts used together don't follow typical accounting patterns (e.g. debiting Revenue against an Expense account), which may indicate a misclassification.",
    },
    "Excessive Amount": {
        "icon": "📈",
        "short": "Abnormally Large Amount",
        "desc": "This entry is significantly larger than the historical average for this account — it stands out as an outlier and warrants further review.",
    },
}

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="Audit Anomaly Explainer",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# Custom CSS
# ---------------------------------------------------------------------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

/* ── Dark theme base ── */
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
#MainMenu, footer {visibility: hidden;}
header[data-testid="stHeader"] {background: transparent !important;}

.stApp {
    background: linear-gradient(160deg, #0a0f1a 0%, #111827 50%, #0d1520 100%);
    color: #ffffff;
}

/* ALL text → pure white, bold, bigger */
.stApp [data-testid="stMarkdownContainer"] p,
.stApp [data-testid="stMarkdownContainer"] li { color: #ffffff !important; font-weight: 500; font-size: 1.08rem !important; }
.stApp [data-testid="stCaptionContainer"] p { color: #e2e8f0 !important; font-weight: 500; }
.stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp h5 { color: #ffffff !important; font-weight: 700 !important; }
.stApp label, .stApp [data-testid="stWidgetLabel"] p { color: #ffffff !important; font-weight: 600 !important; }
.stApp [data-testid="stExpander"] { background: #1a2332; border: 1px solid #2a3a4e; border-radius: 10px; }
.stApp [data-testid="stExpanderDetails"] { background: #1a2332; }
.stApp [data-testid="stExpander"] summary p { color: #ffffff !important; font-weight: 600 !important; font-size: 1.05rem !important; }

/* Hero banner */
.hero {
    background: linear-gradient(135deg, #0c1929 0%, #162a3e 40%, #1a3a5c 100%);
    border: 1px solid #1e3a5f;
    border-radius: 16px;
    padding: 2.5rem 3rem;
    color: white;
    margin-bottom: 2rem;
    position: relative;
    overflow: hidden;
    box-shadow: 0 4px 24px rgba(0,0,0,0.4);
}
.hero::before {
    content: '';
    position: absolute;
    top: -50%; right: -20%;
    width: 500px; height: 500px;
    background: radial-gradient(circle, rgba(56,189,248,0.08) 0%, transparent 70%);
    border-radius: 50%;
}
.hero::after {
    content: '📊';
    position: absolute;
    right: 3rem; top: 50%;
    transform: translateY(-50%);
    font-size: 5rem;
    opacity: 0.12;
}
.hero h1 { color: white !important; margin: 0 0 0.3rem 0; font-size: 2.2rem; font-weight: 700; }
.hero p { color: #7eb8da; margin: 0; font-size: 1.05rem; }

/* Step headers */
.step-header {
    display: flex;
    align-items: center;
    gap: 0.75rem;
    margin: 2rem 0 1rem 0;
}
.step-num {
    background: linear-gradient(135deg, #2563eb, #1d4ed8);
    color: white;
    width: 32px; height: 32px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-weight: 700;
    font-size: 0.9rem;
    flex-shrink: 0;
    box-shadow: 0 2px 8px rgba(37,99,235,0.4);
}
.step-title { font-size: 1.35rem; font-weight: 700; color: #ffffff; margin: 0; }

/* Metric cards */
.metric-row { display: flex; gap: 1rem; margin: 1rem 0; }
.m-card {
    flex: 1;
    background: linear-gradient(145deg, #1a2332, #1e2a3a);
    border: 1px solid #2a3a4e;
    border-radius: 12px;
    padding: 1.2rem;
    text-align: center;
    box-shadow: 0 2px 12px rgba(0,0,0,0.3);
}
.m-card .m-val { font-size: 1.8rem; font-weight: 700; color: #60a5fa; margin: 0; }
.m-card .m-label { font-size: 0.85rem; color: #cbd5e1; font-weight: 600; margin: 0.2rem 0 0 0; text-transform: uppercase; letter-spacing: 0.5px; }
.m-card-warn .m-val { color: #f59e0b; }
.m-card-danger { background: linear-gradient(145deg, #7f1d1d, #991b1b) !important; border-color: #dc2626 !important; }
.m-card-danger .m-val { color: #ffffff; }
.m-card-danger .m-label { color: #fecaca; }
.m-card-ok .m-val { color: #34d399; }

/* Flag badge */
.flag-badge {
    display: inline-block;
    padding: 0.35rem 0.9rem;
    border-radius: 20px;
    font-size: 0.95rem;
    font-weight: 600;
    margin: 0.15rem 0.2rem;
}
.flag-high { background: rgba(239,68,68,0.15); color: #fca5a5; border: 1px solid rgba(239,68,68,0.3); }
.flag-medium { background: rgba(245,158,11,0.15); color: #fcd34d; border: 1px solid rgba(245,158,11,0.3); }
.flag-dup { background: rgba(96,165,250,0.15); color: #93c5fd; border: 1px solid rgba(96,165,250,0.3); }

/* Observation box */
.obs-box {
    background: linear-gradient(to right, #1a2a3e, #1e2d40);
    border-left: 4px solid #2563eb;
    padding: 1.25rem 1.5rem;
    border-radius: 0 10px 10px 0;
    margin: 0.75rem 0;
    line-height: 1.8;
    font-size: 1.05rem;
    font-weight: 500;
    color: #ffffff;
}

/* Detail table */
.detail-grid {
    display: grid;
    grid-template-columns: auto 1fr;
    gap: 0.5rem 1.2rem;
    font-size: 1.05rem;
}
.detail-label { color: #93c5fd; font-weight: 700; }
.detail-value { color: #ffffff; font-weight: 500; }

/* Status pill */
.pill {
    display: inline-block;
    padding: 0.2rem 0.7rem;
    border-radius: 12px;
    font-size: 0.78rem;
    font-weight: 600;
}
.pill-accepted { background: rgba(52,211,153,0.25); color: #ffffff; }
.pill-revision { background: rgba(239,68,68,0.25); color: #ffffff; }
.pill-pending { background: rgba(148,163,184,0.25); color: #ffffff; }

/* Disclaimer */
.disclaimer-box {
    background: linear-gradient(to right, rgba(245,158,11,0.1), rgba(30,41,59,0.5));
    border: 1px solid rgba(245,158,11,0.25);
    border-left: 4px solid #f59e0b;
    border-radius: 0 10px 10px 0;
    padding: 1rem 1.25rem;
    font-size: 1rem;
    font-weight: 600;
    color: #ffffff;
    margin-top: 2.5rem;
}

/* Sidebar toggle — make expand/collapse arrow visible on dark bg */
[data-testid="stSidebarCollapsedControl"] {
    background: #1e293b !important;
    border: 1px solid #3b82f6 !important;
    border-radius: 8px !important;
    padding: 6px !important;
}
/* Catch-all: ANY button used to open/close sidebar */
[data-testid="stSidebarCollapsedControl"] button,
[data-testid="stSidebarCollapsedControl"] > *,
[data-testid="collapsedControl"] button,
[data-testid="collapsedControl"] > *,
[data-testid="stSidebar"] button[kind="headerNoPadding"],
button[aria-label*="idebar"] {
    color: #ffffff !important;
    background: #1e293b !important;
    border: none !important;
    opacity: 1 !important;
    visibility: visible !important;
}
[data-testid="stSidebarCollapsedControl"] svg,
[data-testid="collapsedControl"] svg,
[data-testid="stSidebar"] button[kind="headerNoPadding"] svg,
button[aria-label*="idebar"] svg {
    fill: #ffffff !important;
    color: #ffffff !important;
    stroke: #ffffff !important;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #111827 0%, #0f1a2b 100%) !important;
    border-right: 1px solid #1e2d40;
}
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p,
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] h3 { color: #ffffff !important; font-weight: 600 !important; }
[data-testid="stSidebar"] [data-testid="stCaptionContainer"] p { color: #e2e8f0 !important; font-weight: 500 !important; font-size: 0.92rem !important; }

/* ── Global dark-mode overrides for ALL Streamlit widgets ── */

/* Larger text globally */
.stApp { font-size: 1.1rem; }
.stApp [data-testid="stCaptionContainer"] p { font-size: 1.02rem !important; line-height: 1.7; color: #e2e8f0 !important; font-weight: 500 !important; }
.stApp [data-testid="stMarkdownContainer"] p { font-size: 1.08rem; }

/* ALL buttons → white text on dark bg */
.stApp button {
    color: #ffffff !important;
    font-weight: 600 !important;
    background: #1e293b !important;
    border: 1px solid #3b82f6 !important;
}
.stApp button:hover { background: #2563eb !important; }
.stApp button[kind="primary"] {
    background: linear-gradient(135deg, #2563eb, #1d4ed8) !important;
    border: none !important;
    box-shadow: 0 2px 12px rgba(37,99,235,0.4);
}

/* File uploader — force ALL inner text white */
.stApp [data-testid="stFileUploader"],
.stApp [data-testid="stFileUploaderDropzone"] {
    background: #1a2332 !important;
    border: 2px dashed #3b82f6 !important;
    border-radius: 10px;
}
.stApp [data-testid="stFileUploader"] *,
.stApp [data-testid="stFileUploaderDropzone"] *,
.stApp [data-testid="stFileUploaderDropzone"] span,
.stApp [data-testid="stFileUploaderDropzone"] p,
.stApp [data-testid="stFileUploaderDropzone"] div,
.stApp [data-testid="stFileUploaderDropzone"] label {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}
.stApp [data-testid="stFileUploader"] small,
.stApp [data-testid="stFileUploaderDropzone"] small {
    color: #94a3b8 !important;
    -webkit-text-fill-color: #94a3b8 !important;
}
.stApp [data-testid="stFileUploader"] button,
.stApp [data-testid="stFileUploaderDropzone"] button {
    background: #2563eb !important;
    color: #ffffff !important;
}

/* Text inputs + password */
.stApp input,
.stApp textarea,
.stApp [data-testid="stTextInput"] input,
.stApp [data-baseweb="input"] input {
    color: #ffffff !important;
    background: #1a2332 !important;
    border-color: #2a3a4e !important;
}
.stApp input::placeholder,
.stApp textarea::placeholder { color: #64748b !important; }

/* Selectbox / dropdown — force white on everything inside */
.stApp [data-baseweb="select"],
.stApp [data-baseweb="select"] > div { background-color: #1a2332 !important; }
.stApp [data-baseweb="select"] span,
.stApp [data-baseweb="select"] div,
.stApp [data-baseweb="select"] p,
.stApp [data-baseweb="select"] input,
.stApp [data-baseweb="select"] svg,
.stApp [data-baseweb="select"] [class*="ValueContainer"],
.stApp [data-baseweb="select"] [class*="singleValue"],
.stApp [data-baseweb="select"] [class*="option"],
.stApp [data-baseweb="select"] [class*="placeholder"] {
    color: #ffffff !important;
    fill: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* Selectbox / dropdown — open menu */
.stApp [data-baseweb="popover"],
.stApp [data-baseweb="menu"],
.stApp [role="listbox"] { background: #1e293b !important; }
.stApp [data-baseweb="popover"] li,
.stApp [data-baseweb="menu"] li,
.stApp [role="option"],
.stApp [role="option"] span,
.stApp [role="option"] div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}
.stApp [role="option"]:hover,
.stApp [role="option"][aria-selected="true"],
.stApp [data-baseweb="menu"] li:hover { background: #2563eb !important; }

/* Multiselect tags */
.stApp [data-baseweb="tag"] { background: #2563eb !important; }
.stApp [data-baseweb="tag"] span { color: #ffffff !important; }

/* Number input / slider */
.stApp [data-testid="stNumberInput"] input { color: #ffffff !important; background: #1a2332 !important; }
.stApp [data-baseweb="slider"] div { color: #ffffff !important; }

/* Date input */
.stApp [data-testid="stDateInput"] input { color: #ffffff !important; background: #1a2332 !important; }

/* Checkbox */
.stApp [data-testid="stCheckbox"] label span { color: #ffffff !important; }

/* Download buttons */
.stApp [data-testid="stDownloadButton"] button {
    background: #1e293b !important;
    border: 1px solid #3b82f6 !important;
}

/* Light-background containers (alerts, toasts) → dark text */
.stApp [data-testid="stAlert"] p,
.stApp [data-testid="stAlert"] span,
.stApp [data-testid="stNotification"] p,
.stApp [role="alert"] p,
.stApp [role="alert"] span { color: #1a1a2e !important; font-weight: 600 !important; }

/* Data tables */
.stApp [data-testid="stDataFrame"] { color: #ffffff !important; }

/* Popover content (Supporting Policies, Baseline) */
.stApp [data-testid="stPopover"] [data-testid="stMarkdownContainer"] p { color: #ffffff !important; }

/* Policy card */
.policy-card {
    background: #1a2a3e;
    border: 1px solid #2a3a4e;
    border-radius: 10px;
    padding: 1rem 1.25rem;
    margin: 0.5rem 0;
}
.policy-card .policy-title { color: #ffffff; font-weight: 700; font-size: 1.05rem; margin-bottom: 0.3rem; }
.policy-card .policy-section { color: #e2e8f0; font-size: 0.95rem; font-weight: 500; margin-bottom: 0.5rem; }
.policy-card .policy-score { display: inline-block; background: rgba(96,165,250,0.2); color: #ffffff; padding: 0.2rem 0.6rem; border-radius: 8px; font-size: 0.9rem; font-weight: 600; }
.policy-card .policy-text { color: #e2e8f0; font-size: 1rem; font-weight: 500; line-height: 1.7; margin-top: 0.5rem; white-space: pre-wrap; }
</style>
""", unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Sidebar
# ---------------------------------------------------------------------------
with st.sidebar:
    st.markdown("### ⚙️ Settings")

    st.markdown("**API Key**")
    google_key = st.text_input(
        "Google API Key",
        type="password",
        value=os.environ.get("GOOGLE_API_KEY", ""),
        placeholder="AIza...",
        label_visibility="collapsed",
    )
    if google_key:
        st.caption("✅ Key provided")
    else:
        st.caption("🔗 [Get a key](https://aistudio.google.com/apikey)")

    st.markdown("---")
    st.markdown("**Model**")
    selected_model = st.selectbox(
        "model", options=list(AVAILABLE_MODELS.keys()),
        format_func=lambda k: AVAILABLE_MODELS[k],
        label_visibility="collapsed",
    )

    st.markdown("---")
    st.markdown("**Detection Settings**")
    close_date = st.date_input("Period close date", value=pd.to_datetime("2026-01-03"))
    st.caption("Entries posted after this date are flagged as post-close adjustments.")

    round_threshold = st.number_input("Round-dollar threshold ($)", value=50000, step=10000)
    st.caption("Flag transactions above this amount if they are exact round numbers (e.g. $50,000). Round amounts may indicate estimates or fabricated entries.")

    excessive_mult = st.slider("Excessive amount multiplier", 2.0, 10.0, 3.0, 0.5)
    st.caption(f"Flag entries whose amount exceeds {excessive_mult:.1f}× the average for that account — potential outliers.")

    top_k = st.slider("Policy chunks to retrieve", 1, 5, 3)
    st.caption("How many policy excerpts to retrieve from the handbook for each finding. More excerpts = richer context but higher cost.")

    run_baseline = st.checkbox("Run baseline comparison", value=False)
    st.caption("Compare AI output with and without policy context to measure RAG improvement.")

    st.markdown("---")
    st.markdown("**Chart of Accounts (Optional)**")
    st.caption("Upload your company's chart of accounts so the system can accurately classify account types for anomaly detection.")
    coa_file = st.file_uploader("Upload COA", type=["csv", "xlsx", "xls"],
                                 label_visibility="collapsed", key="coa_upload")
    if coa_file is not None:
        try:
            if coa_file.name.endswith((".xlsx", ".xls")):
                coa_df = pd.read_excel(coa_file)
            else:
                coa_df = pd.read_csv(coa_file)
            st.session_state.account_map = parse_coa(coa_df)
            n_mapped = sum(1 for v in st.session_state.account_map.values() if v != "unknown")
            n_total = len(st.session_state.account_map)
            st.success(f"✅ {n_mapped}/{n_total} accounts classified")
            with st.expander("Preview mapping", expanded=False):
                preview_items = list(st.session_state.account_map.items())[:15]
                for code, cat in preview_items:
                    emoji = {"asset": "🏦", "liability": "📋", "equity": "💰",
                             "revenue": "📈", "expense": "💸"}.get(cat, "❓")
                    st.caption(f"{emoji} {code} → {cat}")
                if len(st.session_state.account_map) > 15:
                    st.caption(f"... and {len(st.session_state.account_map) - 15} more")
        except Exception as e:
            st.error(f"Could not parse COA: {e}")
    elif "account_map" not in st.session_state:
        st.session_state.account_map = None
        st.caption("No COA uploaded — using standard 4-digit convention (1xxx=Asset, 2xxx=Liability, etc.)")

    st.markdown("---")
    st.markdown("**Policy Index**")
    force_rebuild = st.button("🔄 Rebuild", use_container_width=True)

    if google_key and (force_rebuild or "policy_index" not in st.session_state):
        with st.spinner("Building index..."):
            try:
                st.session_state.policy_index = build_or_load_index(
                    str(POLICY_DIR), str(INDEX_DIR), google_key, force_rebuild=force_rebuild
                )
                st.success(f"{len(st.session_state.policy_index.chunks)} chunks")
            except Exception as e:
                st.error(str(e)[:100])

    if "policy_index" in st.session_state:
        st.caption(f"✅ {len(st.session_state.policy_index.chunks)} policy chunks indexed")


# ---------------------------------------------------------------------------
# Hero
# ---------------------------------------------------------------------------
st.markdown("""
<div class="hero">
    <div style="display:flex; align-items:center; gap:1.2rem; margin-bottom:0.5rem;">
        <svg width="52" height="52" viewBox="0 0 52 52" fill="none" xmlns="http://www.w3.org/2000/svg">
            <rect x="2" y="2" width="48" height="48" rx="12" fill="#1e3a5f" stroke="#60a5fa" stroke-width="2"/>
            <path d="M14 38V14h6l4 10 4-10h6v24h-5V23l-5 11-5-11v15z" fill="#60a5fa" opacity="0.15"/>
            <rect x="13" y="14" width="26" height="2" rx="1" fill="#60a5fa"/>
            <rect x="13" y="20" width="20" height="2" rx="1" fill="#60a5fa" opacity="0.7"/>
            <rect x="13" y="26" width="23" height="2" rx="1" fill="#60a5fa" opacity="0.5"/>
            <rect x="13" y="32" width="18" height="2" rx="1" fill="#60a5fa" opacity="0.4"/>
            <circle cx="37" cy="35" r="8" fill="#0c1929" stroke="#60a5fa" stroke-width="2"/>
            <path d="M34 35l2 2 4-4" stroke="#34d399" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>
        <div>
            <h1>Audit Anomaly Explainer</h1>
            <p>Upload any journal entry file → Auto-detect anomalies → Generate policy-grounded audit observations</p>
        </div>
    </div>
    <p style="font-size:0.78rem; color:#4a7a9e; margin-top:0.6rem;">Developed by Xiangyu Yue · Gen AI Project</p>
</div>
""", unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Step 1: Upload
# ---------------------------------------------------------------------------
st.markdown("""
<div class="step-header">
    <div class="step-num">1</div>
    <p class="step-title">Upload Journal Entry Data</p>
</div>
""", unsafe_allow_html=True)

st.caption("Upload any GL export or journal entry CSV. The system will auto-detect anomalies.")

col_up, col_sample = st.columns([5, 1])
with col_up:
    uploaded_file = st.file_uploader("upload", type=["csv", "xlsx", "xls"], label_visibility="collapsed")
with col_sample:
    st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
    use_sample = st.button("📋 Demo Data", use_container_width=True,
                           help="Load the built-in sample journal with 50 entries")

# Persist loaded data in session_state so it survives reruns
if uploaded_file is not None:
    if uploaded_file.name.endswith((".xlsx", ".xls")):
        st.session_state.raw_input_df = pd.read_excel(uploaded_file)
    else:
        st.session_state.raw_input_df = pd.read_csv(uploaded_file)
elif use_sample:
    sample_path = DATA_DIR / "full_journal.csv"
    if sample_path.exists():
        st.session_state.raw_input_df = pd.read_csv(sample_path)

raw_df = st.session_state.get("raw_input_df", None)

if raw_df is not None:
    with st.expander(f"📄 Raw data preview  ·  {len(raw_df)} entries", expanded=False):
        st.dataframe(raw_df, height=250)


# ---------------------------------------------------------------------------
# Step 2: Detect
# ---------------------------------------------------------------------------
if raw_df is not None:
    st.markdown("""
    <div class="step-header">
        <div class="step-num">2</div>
        <p class="step-title">Anomaly Detection</p>
    </div>
    """, unsafe_allow_html=True)

    if st.button("🔍 Run Anomaly Detection", type="primary", use_container_width=False):
        with st.spinner("Scanning journal entries..."):
            flagged_df = run_all_detections(
                raw_df,
                close_date=str(close_date),
                round_threshold=round_threshold,
                excessive_multiplier=excessive_mult,
                account_map=st.session_state.get("account_map"),
            )
            st.session_state.flagged_df = flagged_df
            st.session_state.raw_df = normalize_columns(raw_df.copy())

    if "flagged_df" in st.session_state:
        flagged_df = st.session_state.flagged_df

        if len(flagged_df) == 0:
            st.info("No anomalies detected. Try adjusting detection thresholds in the sidebar.")
        else:
            # Metrics
            n_flagged = len(flagged_df)
            n_high = len(flagged_df[flagged_df["severity"] == "high"]) if "severity" in flagged_df.columns else 0
            total_amt = flagged_df["amount"].astype(float).sum() if "amount" in flagged_df.columns else 0
            flag_types = flagged_df["flag_type"].nunique() if "flag_type" in flagged_df.columns else 0

            n_medium = n_flagged - n_high
            # SVG donut for severity breakdown
            if n_flagged > 0:
                high_pct = n_high / n_flagged
                med_pct = n_medium / n_flagged
                # SVG circle math: circumference = 2*pi*40 ≈ 251.3
                circ = 251.3
                high_len = high_pct * circ
                med_len = med_pct * circ
                med_offset = high_len  # medium starts after high
                donut_svg = f"""
                <div style="display:flex; align-items:center; gap:2rem; margin:1rem 0 0.5rem 0;">
                    <svg width="110" height="110" viewBox="0 0 100 100">
                        <circle cx="50" cy="50" r="40" fill="none" stroke="#1e293b" stroke-width="14"/>
                        <circle cx="50" cy="50" r="40" fill="none" stroke="#ef4444" stroke-width="14"
                                stroke-dasharray="{high_len} {circ - high_len}"
                                stroke-dashoffset="0" transform="rotate(-90 50 50)"/>
                        <circle cx="50" cy="50" r="40" fill="none" stroke="#f59e0b" stroke-width="14"
                                stroke-dasharray="{med_len} {circ - med_len}"
                                stroke-dashoffset="-{med_offset}" transform="rotate(-90 50 50)"/>
                        <text x="50" y="46" text-anchor="middle" fill="#ffffff" font-size="18" font-weight="700">{n_flagged}</text>
                        <text x="50" y="62" text-anchor="middle" fill="#94a3b8" font-size="9" font-weight="600">TOTAL</text>
                    </svg>
                    <div>
                        <div style="display:flex; align-items:center; gap:0.5rem; margin-bottom:0.4rem;">
                            <span style="display:inline-block; width:12px; height:12px; border-radius:3px; background:#ef4444;"></span>
                            <span style="color:#ffffff; font-weight:600; font-size:0.95rem;">High — {n_high} ({high_pct:.0%})</span>
                        </div>
                        <div style="display:flex; align-items:center; gap:0.5rem;">
                            <span style="display:inline-block; width:12px; height:12px; border-radius:3px; background:#f59e0b;"></span>
                            <span style="color:#ffffff; font-weight:600; font-size:0.95rem;">Medium — {n_medium} ({med_pct:.0%})</span>
                        </div>
                    </div>
                </div>
                """
            else:
                donut_svg = ""

            st.markdown(f"""
            <div class="metric-row">
                <div class="m-card m-card-warn"><p class="m-val">{n_flagged}</p><p class="m-label">Anomalies Found</p></div>
                <div class="m-card m-card-danger"><p class="m-val">{n_high}</p><p class="m-label">High Severity</p></div>
                <div class="m-card"><p class="m-val">&#36;{total_amt:,.0f}</p><p class="m-label">Total Flagged Amount</p></div>
                <div class="m-card"><p class="m-val">{flag_types}</p><p class="m-label">Anomaly Types</p></div>
            </div>
            """, unsafe_allow_html=True)

            if donut_svg:
                st.markdown(donut_svg, unsafe_allow_html=True)

            # Anomaly type breakdown — horizontal bar chart
            if "flag_type" in flagged_df.columns:
                # Split combined flag types (e.g. "Excessive Amount | Round-Dollar Amount")
                # and count each base type separately to avoid duplicate bars
                from collections import Counter
                _type_counter = Counter()
                for ft in flagged_df["flag_type"]:
                    for part in str(ft).split(" | "):
                        part = part.strip()
                        if part:
                            _type_counter[part] += 1
                type_counts = pd.Series(_type_counter).sort_values(ascending=False)
                bar_colors = {
                    "Duplicate Entry": "#3b82f6",
                    "Post-Close Entry": "#ef4444",
                    "Round-Dollar Amount": "#f59e0b",
                    "Unusual Account Pairing": "#a855f7",
                    "Excessive Amount": "#ec4899",
                }
                max_count = type_counts.max()

                st.markdown("<div style='margin:1rem 0 1.5rem 0;'>", unsafe_allow_html=True)
                for ftype, count in type_counts.items():
                    info = FLAG_EXPLANATIONS.get(ftype.split(" | ")[0], {})
                    icon = info.get("icon", "⚠️")
                    short = info.get("short", ftype)
                    color = bar_colors.get(ftype.split(" | ")[0], "#6b7280")
                    pct = (count / max_count) * 100
                    st.markdown(f"""
                    <div style="margin-bottom:0.6rem;">
                        <div style="display:flex; justify-content:space-between; margin-bottom:0.25rem;">
                            <span style="color:#ffffff; font-weight:600; font-size:1rem;">{icon} {short}</span>
                            <span style="color:#ffffff; font-weight:700; font-size:1rem;">{count}</span>
                        </div>
                        <div style="background:#1e293b; border-radius:8px; height:28px; overflow:hidden;">
                            <div style="background:{color}; width:{pct}%; height:100%; border-radius:8px; min-width:24px;
                                        display:flex; align-items:center; justify-content:center;
                                        font-size:0.85rem; font-weight:700; color:white;">
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

                # Expandable legend explaining each type
                with st.expander("ℹ️ What do these anomaly types mean?", expanded=False):
                    for ftype in type_counts.index:
                        info = FLAG_EXPLANATIONS.get(ftype.split(" | ")[0], {})
                        icon = info.get("icon", "⚠️")
                        short = info.get("short", ftype)
                        desc = info.get("desc", "")
                        st.markdown(f"**{icon} {short}**")
                        st.caption(desc)
                        st.markdown("")

            with st.expander("📊 View all flagged entries", expanded=False):
                display_cols = [c for c in ["entry_id", "date", "account_debit", "account_credit",
                                            "amount", "posting_user", "flag_type", "severity", "flag_reason"]
                                if c in flagged_df.columns]
                st.dataframe(flagged_df[display_cols], height=300)


# ---------------------------------------------------------------------------
# Step 3: Generate
# ---------------------------------------------------------------------------
if "flagged_df" in st.session_state and len(st.session_state.flagged_df) > 0:
    flagged_df = st.session_state.flagged_df

    st.markdown("""
    <div class="step-header">
        <div class="step-num">3</div>
        <p class="step-title">Generate Audit Observations</p>
    </div>
    """, unsafe_allow_html=True)

    ready = google_key and "policy_index" in st.session_state
    if not ready:
        st.warning("Enter your Google API Key and wait for the policy index to build.")
    else:
        # Filter by anomaly type
        if "flag_type" in flagged_df.columns:
            # Extract unique base types (handle combined types like "Excessive Amount | Round-Dollar Amount")
            all_types_raw = flagged_df["flag_type"].unique().tolist()
            base_types = sorted(set(t.strip() for raw in all_types_raw for t in raw.split(" | ")))
            type_labels = {t: f"{FLAG_EXPLANATIONS.get(t, {}).get('icon', '⚠️')} {FLAG_EXPLANATIONS.get(t, {}).get('short', t)}" for t in base_types}
            selected_types = st.multiselect(
                "Filter by anomaly type (empty = all types)",
                options=base_types,
                format_func=lambda t: type_labels.get(t, t),
            )
        else:
            selected_types = []

        gcol1, gcol2 = st.columns([3, 1])
        with gcol1:
            # Filter entries based on selected types
            if selected_types:
                type_filtered = flagged_df[flagged_df["flag_type"].apply(
                    lambda ft: any(t in ft for t in selected_types)
                )]
            else:
                type_filtered = flagged_df
            all_ids = type_filtered["entry_id"].tolist() if "entry_id" in type_filtered.columns else []
            selected_ids = st.multiselect("Select specific entries (empty = all)", options=all_ids)
        with gcol2:
            st.markdown("<div style='height:1.8rem'></div>", unsafe_allow_html=True)
            gen_clicked = st.button("⚡ Generate All", type="primary", use_container_width=True)

        if gen_clicked:
            if selected_ids:
                entries = flagged_df[flagged_df["entry_id"].isin(selected_ids)].to_dict("records")
            elif selected_types:
                entries = type_filtered.to_dict("records")
            else:
                entries = flagged_df.to_dict("records")

            results = []
            baseline_results = []
            progress = st.progress(0)
            status_text = st.empty()
            _gen_start_time = time.time()

            for i, anomaly in enumerate(entries):
                eid = anomaly.get("entry_id", i + 1)
                # Calculate ETA
                elapsed = time.time() - _gen_start_time
                if i > 0:
                    avg_per_entry = elapsed / i
                    remaining = avg_per_entry * (len(entries) - i)
                    if remaining >= 60:
                        eta_str = f"~{remaining/60:.1f} min remaining"
                    else:
                        eta_str = f"~{remaining:.0f}s remaining"
                else:
                    eta_str = "estimating..."
                status_text.caption(f"Processing {eid}  ({i+1}/{len(entries)}) — {eta_str}")
                progress.progress((i + 1) / len(entries))

                # Retrieve policies with 429 handling
                try:
                    query = f"{anomaly.get('flag_reason', '')} {anomaly.get('account_debit', '')} {anomaly.get('account_credit', '')}"
                    retrieved = st.session_state.policy_index.search(query, google_key, top_k=top_k)
                except Exception as e:
                    if "429" in str(e) or "ResourceExhausted" in str(e):
                        status_text.caption(f"⏳ Rate limited on {eid} — waiting 30s then retrying...")
                        time.sleep(30)
                        try:
                            retrieved = st.session_state.policy_index.search(query, google_key, top_k=top_k)
                        except Exception:
                            retrieved = []
                    else:
                        retrieved = []

                # Generate observation
                try:
                    result = generate_observation(anomaly, retrieved, google_key, model=selected_model)
                    result["retrieved_policies"] = retrieved
                    results.append(result)
                except Exception as e:
                    if "429" in str(e) or "ResourceExhausted" in str(e):
                        status_text.caption(f"⏳ Rate limited on {eid} — waiting 30s then retrying...")
                        time.sleep(30)
                        try:
                            result = generate_observation(anomaly, retrieved, google_key, model=selected_model)
                            result["retrieved_policies"] = retrieved
                            results.append(result)
                        except Exception as e2:
                            results.append({
                                "entry_id": eid, "observation": f"API rate limited. Try again in a few minutes.",
                                "had_context": True, "retrieved_policies": retrieved, "error": True,
                            })
                    else:
                        results.append({
                            "entry_id": eid, "observation": f"Error: {e}",
                            "had_context": True, "retrieved_policies": retrieved, "error": True,
                        })

                if run_baseline:
                    try:
                        bl = generate_observation_baseline(anomaly, google_key, model=selected_model)
                        baseline_results.append(bl)
                    except Exception as e:
                        baseline_results.append({"entry_id": eid, "observation": f"Error: {e}", "had_context": False})

                time.sleep(4)  # Pace requests for Gemini free-tier

            progress.empty()
            status_text.empty()
            st.session_state.results = results
            st.session_state.baseline_results = baseline_results
            st.session_state.processed_entries = entries
            st.success(f"✅ Generated {len(results)} observations")


# ---------------------------------------------------------------------------
# Step 4: Review
# ---------------------------------------------------------------------------
if "results" in st.session_state and st.session_state.results:
    results = st.session_state.results
    baseline_results = st.session_state.baseline_results
    entries = st.session_state.processed_entries

    st.markdown("""
    <div class="step-header">
        <div class="step-num">4</div>
        <p class="step-title">Review & Export</p>
    </div>
    """, unsafe_allow_html=True)

    if "accepted" not in st.session_state:
        st.session_state.accepted = {}

    n_acc = sum(1 for v in st.session_state.accepted.values() if v is True)
    n_rej = sum(1 for v in st.session_state.accepted.values() if v is False)
    n_pen = len(results) - n_acc - n_rej

    st.markdown(f"""
    <div class="metric-row">
        <div class="m-card"><p class="m-val">{len(results)}</p><p class="m-label">Total</p></div>
        <div class="m-card m-card-ok"><p class="m-val">{n_acc}</p><p class="m-label">Accepted</p></div>
        <div class="m-card m-card-danger"><p class="m-val">{n_rej}</p><p class="m-label">Revision</p></div>
        <div class="m-card"><p class="m-val">{n_pen}</p><p class="m-label">Pending</p></div>
    </div>
    """, unsafe_allow_html=True)

    for i, (result, anomaly) in enumerate(zip(results, entries)):
        eid = result.get("entry_id", f"Entry {i+1}")
        flag = anomaly.get("flag_reason", "Unknown")
        flag_type = anomaly.get("flag_type", "")
        severity = anomaly.get("severity", "medium")
        related = anomaly.get("related_entries", [])
        status = st.session_state.accepted.get(eid)
        is_error = result.get("error", False)

        # Build header
        if status is True:
            pill = '<span class="pill pill-accepted">✅ Accepted</span>'
        elif status is False:
            pill = '<span class="pill pill-revision">🔄 Revision</span>'
        else:
            pill = '<span class="pill pill-pending">⏳ Pending</span>'

        sev_class = "high" if severity == "high" else "medium"

        # Build a plain-language summary for the expander title
        flag_info = FLAG_EXPLANATIONS.get(flag_type.split(" | ")[0] if flag_type else "", {})
        flag_short = flag_info.get("short", flag_type)
        flag_icon = flag_info.get("icon", "⚠️")
        flag_desc_text = flag_info.get("desc", "")

        with st.expander(f"{flag_icon} {eid}  ·  {flag_short}  ·  ${float(anomaly.get('amount', 0)):,.0f}", expanded=(i == 0)):
            st.markdown(f'{pill} <span class="flag-badge flag-{sev_class}">{severity.upper()}</span>', unsafe_allow_html=True)

            col_d, col_o = st.columns([2, 3])

            with col_d:
                st.markdown("##### 📋 Entry Details")
                # Format date — strip time component
                _raw_date = anomaly.get('date', 'N/A')
                try:
                    _display_date = pd.to_datetime(_raw_date).strftime('%Y-%m-%d')
                except Exception:
                    _display_date = str(_raw_date).split(' ')[0] if ' ' in str(_raw_date) else str(_raw_date)
                st.markdown(f"""
                <div class="detail-grid">
                    <span class="detail-label">Date</span><span class="detail-value">{_display_date}</span>
                    <span class="detail-label">Debit</span><span class="detail-value">{anomaly.get('account_debit', 'N/A')}</span>
                    <span class="detail-label">Credit</span><span class="detail-value">{anomaly.get('account_credit', 'N/A')}</span>
                    <span class="detail-label">Amount</span><span class="detail-value"><strong>${float(anomaly.get('amount', 0)):,.2f}</strong></span>
                    <span class="detail-label">User</span><span class="detail-value">{anomaly.get('posting_user', 'N/A')}</span>
                    <span class="detail-label">Description</span><span class="detail-value">{anomaly.get('description', 'N/A')}</span>
                </div>
                """, unsafe_allow_html=True)

                # Show the plain-language "Why flagged" box
                st.markdown("")
                st.markdown(f"**{flag_icon} Why flagged: {flag_short}**")
                if flag_desc_text:
                    st.caption(flag_desc_text)
                # Escape $ to prevent LaTeX rendering
                flag_escaped = str(flag).replace("$", "&#36;")
                st.markdown(f"**Detail:** <span style='color:#e2e8f0;'>{flag_escaped}</span>", unsafe_allow_html=True)

                # Show related entries for duplicates
                if related:
                    st.markdown(f"**🔗 Related entries:** {', '.join(str(r) for r in related)}")
                    # Show the matching entry data
                    if "raw_df" in st.session_state:
                        raw = st.session_state.raw_df
                        if "entry_id" in raw.columns:
                            matches = raw[raw["entry_id"].isin(related)]
                            if len(matches) > 0:
                                st.markdown("**Matching entry:**")
                                for _, match_row in matches.iterrows():
                                    st.caption(
                                        f"{match_row.get('entry_id', '')} · "
                                        f"{match_row.get('date', '')} · "
                                        f"${float(match_row.get('amount', 0)):,.2f} · "
                                        f"{match_row.get('description', '')}"
                                    )

                if "input_tokens" in result:
                    st.caption(f"📊 Tokens: {result.get('input_tokens', 0)} in / {result.get('output_tokens', 0)} out")

            with col_o:
                st.markdown("##### 📝 Audit Observation")
                if is_error:
                    st.error(result["observation"])
                else:
                    # Escape $ signs to prevent LaTeX rendering in Streamlit
                    obs_html = result["observation"].replace("$", "&#36;")
                    # Convert markdown bold to HTML bold for proper display
                    import re as _re
                    obs_html = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", obs_html)
                    # Color-code observation section headings
                    _section_colors = {
                        "Condition:": "#60a5fa",   # blue
                        "Criteria:": "#a78bfa",    # purple
                        "Cause:": "#f59e0b",       # amber
                        "Effect:": "#ef4444",       # red
                        "Recommendation:": "#34d399", # green
                    }
                    for _sec, _clr in _section_colors.items():
                        obs_html = obs_html.replace(
                            f"<strong>{_sec}</strong>",
                            f'<strong style="color:{_clr}; font-size:1.08rem;">{_sec}</strong>'
                        )
                    # Convert newlines to <br>
                    obs_html = obs_html.replace("\n\n", "<br><br>").replace("\n", "<br>")
                    st.markdown(f'<div class="obs-box">{obs_html}</div>', unsafe_allow_html=True)

                    # Observation completeness check
                    _obs_text = result["observation"]
                    _required = ["Condition:", "Criteria:", "Cause:", "Effect:", "Recommendation:"]
                    _missing = [s.replace(":", "") for s in _required if s not in _obs_text]
                    if _missing:
                        st.markdown(
                            f'<div style="background:rgba(239,68,68,0.12); border:1px solid rgba(239,68,68,0.3); '
                            f'border-radius:8px; padding:0.5rem 0.8rem; margin-top:0.4rem; font-size:0.92rem;">'
                            f'⚠️ <strong style="color:#fca5a5;">Incomplete:</strong> '
                            f'<span style="color:#e2e8f0;">Missing {", ".join(_missing)}. '
                            f'Try regenerating — may be due to API rate limiting.</span></div>',
                            unsafe_allow_html=True,
                        )

                if baseline_results and i < len(baseline_results):
                    with st.popover("📊 Baseline Comparison"):
                        st.markdown("**Without policy context:**")
                        baseline_text = baseline_results[i].get("observation", "")
                        baseline_html = baseline_text.replace("$", "&#36;")
                        import re as _re2
                        baseline_html = _re2.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", baseline_html)
                        for _sec, _clr in _section_colors.items():
                            baseline_html = baseline_html.replace(
                                f"<strong>{_sec}</strong>",
                                f'<strong style="color:{_clr}; font-size:1.08rem;">{_sec}</strong>'
                            )
                        baseline_html = baseline_html.replace("\n\n", "<br><br>").replace("\n", "<br>")
                        st.markdown(f'<div class="obs-box">{baseline_html}</div>', unsafe_allow_html=True)

            if result.get("retrieved_policies"):
                with st.popover("📚 Supporting Policies"):
                    st.markdown("**Policies retrieved from the company handbook that relate to this finding:**")
                    st.markdown("")
                    for j, chunk in enumerate(result["retrieved_policies"]):
                        title = chunk.get("doc_title", "Unknown Policy")
                        section = chunk.get("section", "")
                        score = chunk.get("score", 0)
                        text = chunk.get("text", "")
                        relevance_pct = f"{score * 100:.0f}%"
                        st.markdown(f"""
                        <div class="policy-card">
                            <div class="policy-title">📄 {title}</div>
                            <div class="policy-section">Section: {section} &nbsp; <span class="policy-score">Relevance: {relevance_pct}</span></div>
                            <div class="policy-text">{text}</div>
                        </div>
                        """, unsafe_allow_html=True)

            b1, b2, _ = st.columns([1, 1, 3])
            with b1:
                if st.button("✅ Accept", key=f"a_{eid}", use_container_width=True):
                    st.session_state.accepted[eid] = True
                    st.rerun()
            with b2:
                if st.button("🔄 Revise", key=f"r_{eid}", use_container_width=True):
                    st.session_state.accepted[eid] = False
                    st.rerun()

    # Export
    st.markdown("---")
    e1, e2, e3 = st.columns(3)
    with e1:
        rows = []
        for res, ano in zip(results, entries):
            s = st.session_state.accepted.get(res.get("entry_id", ""))
            rows.append({
                "entry_id": res.get("entry_id", ""),
                "date": ano.get("date", ""),
                "account_debit": ano.get("account_debit", ""),
                "account_credit": ano.get("account_credit", ""),
                "amount": ano.get("amount", ""),
                "flag_type": ano.get("flag_type", ""),
                "flag_reason": ano.get("flag_reason", ""),
                "observation": res.get("observation", ""),
                "status": "Accepted" if s is True else ("Revision" if s is False else "Pending"),
            })
        st.download_button("📥 Export CSV", pd.DataFrame(rows).to_csv(index=False),
                           "audit_observations.csv", "text/csv", use_container_width=True)

    with e2:
        md = ["# Audit Observations Report\n", f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}\n"]
        for res, ano in zip(results, entries):
            s = st.session_state.accepted.get(res.get("entry_id", ""))
            label = "Accepted" if s is True else ("Revision" if s is False else "Pending")
            md.append(f"\n## {res.get('entry_id', '')} — {label}")
            md.append(f"**Type:** {ano.get('flag_type', '')}  |  **Amount:** ${float(ano.get('amount', 0)):,.2f}")
            md.append(f"**Flag:** {ano.get('flag_reason', '')}\n")
            md.append(res.get("observation", ""))
            md.append("\n---")
        st.download_button("📥 Export Markdown", "\n".join(md),
                           "audit_observations.md", "text/markdown", use_container_width=True)

    with e3:
        # DOCX export
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io as _io

            doc = Document()
            # Title
            title_p = doc.add_heading("Audit Observations Report", level=0)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("")

            for res, ano in zip(results, entries):
                s = st.session_state.accepted.get(res.get("entry_id", ""))
                label = "Accepted" if s is True else ("Revision" if s is False else "Pending")

                # Entry heading
                doc.add_heading(f"{res.get('entry_id', '')} — {label}", level=2)

                # Details table
                table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
                details = [
                    ("Type", ano.get("flag_type", "")),
                    ("Amount", f"${float(ano.get('amount', 0)):,.2f}"),
                    ("Accounts", f"{ano.get('account_debit', '')} → {ano.get('account_credit', '')}"),
                    ("Flag Reason", ano.get("flag_reason", "")),
                    ("Status", label),
                ]
                for row_idx, (lbl, val) in enumerate(details):
                    table.rows[row_idx].cells[0].text = lbl
                    table.rows[row_idx].cells[1].text = str(val)
                    for cell in table.rows[row_idx].cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style.font.size = Pt(10)

                doc.add_paragraph("")

                # Observation text
                obs_text = res.get("observation", "")
                obs_para = doc.add_paragraph()
                # Parse sections and bold the headings
                import re as _re_docx
                parts = _re_docx.split(r"(\*\*(?:Condition|Criteria|Cause|Effect|Recommendation):\*\*)", obs_text)
                for part in parts:
                    heading_match = _re_docx.match(r"\*\*(.+?)\*\*", part)
                    if heading_match:
                        run = obs_para.add_run(heading_match.group(1) + " ")
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
                    else:
                        run = obs_para.add_run(part)
                        run.font.size = Pt(10)

                doc.add_paragraph("─" * 50)

            # Disclaimer
            disc = doc.add_paragraph()
            disc_run = disc.add_run(
                "⚠️ AI-Generated Draft — Requires Auditor Review. "
                "Observations are generated by Gemini with RAG-retrieved policy context. "
                "Always verify citations against source documents."
            )
            disc_run.italic = True
            disc_run.font.size = Pt(9)
            disc_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            docx_buf = _io.BytesIO()
            doc.save(docx_buf)
            docx_buf.seek(0)
            st.download_button("📥 Export Word", docx_buf.getvalue(),
                               "audit_observations.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True)
        except ImportError:
            st.caption("Install `python-docx` for Word export")


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
st.markdown("""
<div class="disclaimer-box">
    ⚠️ <strong>AI-Generated Draft — Requires Auditor Review</strong><br>
    Observations are generated by Gemini with RAG-retrieved policy context.
    Always verify citations against source documents. This tool assists with drafting — it does not replace professional audit judgment.
</div>
""", unsafe_allow_html=True)
